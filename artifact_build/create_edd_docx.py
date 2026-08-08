from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'deliverables' / 'part1'
LOGO = ROOT / 'assets' / 'branding' / 'lifeready-training-logo.png'
EVIDENCE = ROOT / 'assets' / 'evidence'
WIREFRAMES = ROOT / 'assets' / 'wireframes'
DOCX_FILE = OUT / 'EDD-v1-LifeReady-Training.docx'

NAVY = '102A43'
TEAL = '047E87'
ORANGE = 'F59E0B'
LIGHT = 'F5F8FA'
MID = 'D9E2EC'
MUTED = '52606D'
WHITE = 'FFFFFF'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in('w:tcW')
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    tr_pr.append(header)


def font(run, size=None, color=NAVY, bold=None, italic=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text(paragraph, text, size=11, color=NAVY, bold=False, italic=False):
    run = paragraph.add_run(text)
    font(run, size, color, bold, italic)
    return run


def add_para(doc, text='', style='Body', align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        add_text(p, text, 11 if style == 'Body' else None)
    return p


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ('Heading 1', 16, '2E74B5', 16, 8),
        ('Heading 2', 13, '2E74B5', 12, 6),
        ('Heading 3', 12, '1F4D78', 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(header, 'LifeReady Training | COMP229 Project Part 1', 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, 'LifeReady Training - External Design Document v1', 8, MUTED)


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    add_text(p, text, {1: 16, 2: 13, 3: 12}[level], {1: '2E74B5', 2: '2E74B5', 3: '1F4D78'}[level], True)
    return p


def body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='Normal')
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, 11, NAVY, True)
        add_text(p, text[len(bold_prefix):], 11)
    else:
        add_text(p, text, 11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    add_text(p, text, 11)
    return p


def make_wireframe(filename, title, rows):
    WIREFRAMES.mkdir(parents=True, exist_ok=True)
    target = WIREFRAMES / filename
    im = Image.new('RGB', (900, 610), 'white')
    draw = ImageDraw.Draw(im)
    try:
        regular = ImageFont.truetype('arial.ttf', 28)
        small = ImageFont.truetype('arial.ttf', 21)
        bold = ImageFont.truetype('arialbd.ttf', 28)
    except OSError:
        regular = small = bold = ImageFont.load_default()
    draw.rounded_rectangle((8, 8, 892, 602), radius=18, outline=f'#{TEAL}', width=4)
    draw.rounded_rectangle((8, 8, 892, 92), radius=18, fill=f'#{NAVY}')
    draw.text((34, 34), title, font=bold, fill='white')
    y = 124
    for label, kind in rows:
        if kind == 'hero':
            draw.rounded_rectangle((34, y, 866, y + 104), radius=10, fill=f'#{LIGHT}')
            draw.text((58, y + 36), label, font=bold, fill=f'#{TEAL}')
            y += 128
        elif kind == 'button':
            draw.rounded_rectangle((34, y, 304, y + 52), radius=9, fill=f'#{ORANGE}')
            draw.text((58, y + 13), label, font=small, fill='white')
            y += 75
        elif kind == 'card':
            draw.rounded_rectangle((34, y, 866, y + 66), radius=8, fill=f'#{LIGHT}', outline=f'#{MID}', width=2)
            draw.text((58, y + 20), label, font=small, fill=f'#{NAVY}')
            y += 90
        else:
            draw.text((42, y), label, font=small, fill=f'#{NAVY}')
            draw.line((34, y + 38, 866, y + 38), fill=f'#{MID}', width=2)
            y += 66
    im.save(target)
    return target


def wireframe_table(doc):
    frames = [
        make_wireframe('home.png', 'Home / Landing Page', [('LifeReady Training', 'hero'), ('Find a course near you', 'normal'), ('View classes', 'button')]),
        make_wireframe('classes.png', 'Training Class List', [('Search and category filter', 'normal'), ('BLS - Toronto - Feb 15', 'card'), ('CPR/AED - Toronto - Feb 22', 'card'), ('View class details', 'button')]),
        make_wireframe('signin.png', 'Sign Up / Sign In', [('Name, email, password', 'normal'), ('Create account', 'button'), ('Sign in', 'button')]),
        make_wireframe('admin.png', 'Admin Class Management', [('Add training class', 'button'), ('Class record - edit / delete', 'card'), ('Capacity, date, price fields', 'normal')]),
    ]
    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [4680, 4680])
    for i, target in enumerate(frames):
        cell = table.cell(i // 2, i % 2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(target), width=Inches(3.0))
    return table


def table_with_rows(doc, rows, widths, header_fill=NAVY):
    table = doc.add_table(rows=0, cols=len(widths))
    set_table_geometry(table, widths)
    for r_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            if r_index == 0:
                set_cell_shading(cell, header_fill)
                add_text(p, value, 8.5, WHITE, True)
            else:
                if r_index % 2 == 0:
                    set_cell_shading(cell, LIGHT)
                add_text(p, value, 8.4, NAVY)
    set_repeat_table_header(table.rows[0])
    return table


def evidence_grid(doc, captures):
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [4680, 4680])
    for start in range(0, len(captures), 2):
        cells = table.add_row().cells
        for index in range(2):
            filename, caption = captures[start + index]
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(EVIDENCE / filename), width=Inches(3.0))
            cap = cell.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(2)
            cap.paragraph_format.space_after = Pt(4)
            add_text(cap, caption, 8, MUTED)
    return table


def page_break(doc):
    doc.add_page_break()


def create_edd():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)

    # Cover: editorial-cover header pattern with project branding.
    doc.add_paragraph().paragraph_format.space_after = Pt(35)
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(str(LOGO), width=Inches(1.65))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    add_text(title, 'External Design Document', 27, NAVY, True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    add_text(subtitle, 'Version 1 - First Release', 14, MUTED)
    meta = doc.add_table(rows=5, cols=2)
    set_table_geometry(meta, [2500, 6860])
    for i, (label, value) in enumerate([
        ('Project', 'LifeReady Training (working name)'),
        ('Course', 'COMP229 - Web Application Development'),
        ('Developer', 'Amir Saad'),
        ('Student ID', '301473849'),
        ('Release focus', 'Database, backend connection, and API testing'),
    ]):
        set_cell_shading(meta.cell(i, 0), LIGHT)
        add_text(meta.cell(i, 0).paragraphs[0], label, 10.5, NAVY, True)
        add_text(meta.cell(i, 1).paragraphs[0], value, 10.5)
    heading(doc, 'Purpose')
    body(doc, 'LifeReady Training is an original MERN application for managing First Aid, CPR/AED, and Basic Life Support (BLS) training classes. This first release establishes the data model, secure Express API, CRUD operations, authentication, authorization, and repeatable API testing.')
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    add_text(note, 'Brand note: the logo was created with OpenAI image generation for this student project. The working name and domain may change before the final release.', 9, MUTED, italic=True)
    page_break(doc)

    heading(doc, 'Table of Contents')
    toc = table_with_rows(doc, [
        ('1', 'Project overview and scope'),
        ('2', 'Architecture and data design'),
        ('3', 'Wireframes for top-level features'),
        ('4', 'API test plan and initial evidence'),
        ('5', 'Part 1 release plan'),
        ('Appendix', 'API and database screenshot evidence'),
    ], [900, 8460], header_fill=LIGHT)
    for cell in toc.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(NAVY)
    heading(doc, '1. Project Overview and Scope')
    body(doc, 'The public-facing site will later provide first-aid information and class discovery. The Part 1 release intentionally prioritizes functional backend components rather than visual polish.')
    heading(doc, 'Primary user roles', 2)
    body(doc, 'Student: creates an account, signs in, and manages their own profile. Administrator: manages training classes and can view user records. This separation demonstrates authorization in addition to authentication.')
    heading(doc, 'First-release goals', 2)
    for item in [
        'Create a MongoDB Atlas database and two collections: Users and TrainingClasses.',
        'Build Node.js and Express MVC APIs.',
        'Test Create, Read, Update, and Delete operations for both user and training-class data.',
        'Restrict class management to administrators using JWT authentication.',
    ]:
        bullet(doc, item)
    page_break(doc)

    heading(doc, '2. Architecture and Data Design')
    body(doc, 'The solution uses a MERN architecture. MongoDB Atlas stores data; Mongoose defines schemas; Express controllers implement MVC behavior; JSON Web Tokens protect authenticated routes. React integration is deliberately reserved for Project Part 2.')
    architecture = [
        ('Layer', 'Implementation', 'Responsibility'),
        ('Database', 'MongoDB Atlas', 'Stores User and TrainingClass documents.'),
        ('Model', 'Mongoose schemas', 'Validation, timestamps, password hashing, class fields.'),
        ('Controller', 'Express MVC controllers', 'Implements CRUD logic and error responses.'),
        ('Routes', 'Express route modules', 'Maps REST endpoints to controller actions.'),
        ('Security', 'JWT + role checks', 'Authenticates users and restricts class management to admins.'),
        ('Testing', 'Thunder Client + repeatable API script', 'Confirms health, auth, authorization, and CRUD responses.'),
    ]
    table_with_rows(doc, architecture, [1600, 2300, 5460])
    heading(doc, 'Collections', 2)
    body(doc, 'User: name, email, password hash, salt, role, timestamps.')
    body(doc, 'TrainingClass: title, category, format, description, class date, duration, location, capacity, price, instructor, status, creator, timestamps.')
    page_break(doc)

    heading(doc, '3. Wireframes - Top-Level Features')
    body(doc, 'These low-fidelity wireframes guide the Part 2 React implementation. They are original planning artifacts, not copied from the reference website.')
    wireframe_table(doc)
    page_break(doc)

    heading(doc, '4. API Test Plan and Initial Evidence')
    body(doc, 'The backend was connected to MongoDB Atlas and verified with both Thunder Client manual requests and a repeatable integration test. The following checks passed in the first release.')
    tests = [
        ('Test', 'Expected result', 'Observed result'),
        ('API health check', '200 OK', 'PASS'),
        ('User create', '201 Created', 'PASS'),
        ('Student authentication', '200 OK + JWT', 'PASS'),
        ('Authorization rule', 'Student blocked from admin class create (403)', 'PASS'),
        ('User read / update / delete', '200 OK', 'PASS'),
        ('Administrator authentication', '200 OK + JWT', 'PASS'),
        ('TrainingClass create / list / update / delete', '201 then 200 responses', 'PASS'),
    ]
    table_with_rows(doc, tests, [3000, 4800, 1560], header_fill=TEAL)
    heading(doc, 'Initial screenshot evidence', 2)
    body(doc, 'The appendix includes saved captures from Thunder Client and MongoDB Atlas. Passwords, JWT values, salts, hashes, and the connection string are not shown.')
    heading(doc, '5. Part 1 Release Plan')
    body(doc, 'The first release delivers a functional API foundation. The next project phase will integrate React pages, navigation, frontend CRUD forms, class browsing, and polished visual design. The final release will add deployment, unit/E2E testing evidence, performance improvements, and CI/CD documentation.')
    page_break(doc)

    evidence_sections = [
        ('Appendix A - Server Startup and Health Evidence', 'These captures show the local API running and responding successfully.', [
            ('vscode-server-running.png', 'Figure A1. VS Code server connected to MongoDB.'),
            ('browser-health-check.png', 'Figure A2. Browser API health response.'),
            ('health-check.png', 'Figure A3. Thunder Client health endpoint: 200 OK.'),
            ('student-create-redacted.png', 'Figure A4. Student account create: 201 Created.'),
        ]),
        ('Appendix B - Authentication and User Data Evidence', 'Sensitive passwords, JWT values, salts, and password hashes are redacted in these captures.', [
            ('student-signin-redacted.png', 'Figure A5. Student sign-in response: 200 OK.'),
            ('admin-signin-redacted.png', 'Figure A6. Administrator sign-in response: 200 OK.'),
            ('atlas-users-redacted.png', 'Figure A7. MongoDB Atlas Users collection with sensitive fields redacted.'),
            ('atlas-empty-trainingclasses.png', 'Figure A8. TrainingClasses collection before a class is created.'),
        ]),
        ('Appendix C - TrainingClass CRUD Evidence', 'The project-specific TrainingClass object is created, read, updated, and deleted through protected API routes.', [
            ('class-create.png', 'Figure A9. TrainingClass create: 201 Created.'),
            ('class-read-list.png', 'Figure A10. TrainingClass list/read: 200 OK.'),
            ('class-update.png', 'Figure A11. TrainingClass update: 200 OK.'),
            ('class-delete.png', 'Figure A12. TrainingClass delete: 200 OK.'),
        ]),
        ('Appendix D - MongoDB Atlas TrainingClass Evidence', 'MongoDB Atlas confirms the TrainingClass document persisted after the create and update operations.', [
            ('atlas-trainingclass.png', 'Figure A13. MongoDB Atlas TrainingClasses collection.'),
            ('atlas-updated-class.png', 'Figure A14. Updated TrainingClass document in Atlas.'),
        ]),
    ]
    for index, (title, description, captures) in enumerate(evidence_sections):
        heading(doc, title)
        body(doc, description)
        evidence_grid(doc, captures)
        if index < len(evidence_sections) - 1:
            page_break(doc)

    doc.core_properties.title = 'LifeReady Training External Design Document v1'
    doc.core_properties.author = 'Amir Saad'
    doc.save(DOCX_FILE)
    print(DOCX_FILE)


if __name__ == '__main__':
    create_edd()
