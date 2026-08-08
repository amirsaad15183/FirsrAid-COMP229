from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "part1" / "LifeReady-Training-Part1-Demo-Video-Script.docx"
LOGO = ROOT / "assets" / "branding" / "lifeready-training-logo.png"

NAVY = "102A43"
TEAL = "047E87"
MUTED = "52606D"


def set_font(run, size, color="000000", bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_page_number(paragraph):
    paragraph.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor.from_string("2E74B5")
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.color.rgb = RGBColor.from_string("2E74B5")
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)


def add_text(doc, text, italic=False, color="000000", size=11):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    set_font(run, size, color)
    run.italic = italic
    return paragraph


def add_cue(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.25)
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, 11, TEAL, True)
    text_run = paragraph.add_run(text)
    set_font(text_run, 11, NAVY)


def add_section(doc, heading, narration, cue=None):
    doc.add_paragraph(heading, style="Heading 1")
    add_text(doc, narration)
    if cue:
        add_cue(doc, "Show", cue)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    setup_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("COMP229 - Web Application Development")
    set_font(header_run, 9, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("LifeReady Training - Project Part 1 | ")
    set_font(footer_run, 9, MUTED)
    add_page_number(footer)

    if LOGO.exists():
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.add_run().add_picture(str(LOGO), width=Inches(1.25))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("Project Part 1 Demo Video Script")
    set_font(title_run, 24, NAVY, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    subtitle_run = subtitle.add_run("LifeReady Training | Amir Saad | Student ID 301473849")
    set_font(subtitle_run, 12, TEAL, True)

    add_text(doc, "Target length: 6 to 8 minutes. Use this as a speaking guide and speak naturally rather than reading every word exactly.", italic=True, color=MUTED)
    add_text(doc, "Important: Hide passwords, JWT tokens, and the MongoDB connection string before you start recording.", italic=True, color="7A5A00")

    add_section(
        doc,
        "0:00 - 0:30 | Introduction",
        "Hello, my name is Amir Saad, student ID 301473849. This is my COMP229 Web Application Development Project Part 1 demo. My project is currently named LifeReady Training. It is an original MERN application for managing First Aid, CPR/AED, and Basic Life Support training classes. I am completing the project independently with professor approval.",
        "Slide 1 of the PowerPoint presentation.",
    )

    add_section(
        doc,
        "0:30 - 1:10 | Part 1 Scope",
        "For Part 1, the application does not need to be visually complete yet. The focus is the database connection, Node and Express backend, MVC structure, authentication, authorization, and CRUD API testing. The React frontend and polished class-browsing pages will be completed in Part 2.",
        "Slide 2, then move to VS Code.",
    )

    add_section(
        doc,
        "1:10 - 1:50 | Database and Collections",
        "I created a MongoDB Atlas project and connected it to my Express application through environment variables. The application uses two main collections. The User collection stores name, email, a salted password hash, role, and timestamps. The TrainingClass collection stores a title, category, class date, location, capacity, price, instructor, status, and the administrator who created the record.",
        "MongoDB Atlas cluster or Data Explorer. Do not show passwords or the connection string.",
    )

    add_section(
        doc,
        "1:50 - 2:40 | MVC Backend Structure",
        "The backend follows MVC. The models define the MongoDB schema and validation. Controllers contain the CRUD logic. Route files map HTTP requests to controller functions. The Express server registers the routes and includes a health check endpoint.",
        "The server/models, server/controllers, and server/routes folders in VS Code.",
    )

    add_section(
        doc,
        "2:40 - 3:30 | Authentication and Authorization",
        "Users can create an account and sign in through the authentication API. A successful sign-in returns a JSON Web Token. The token protects profile routes. The project also uses roles: regular users are students by default, while administrators can manage training classes. The code prevents a public sign-up request from creating an administrator role.",
        "auth.controller.js, including token creation and requireAdmin. Briefly point out the code comments.",
    )

    doc.add_paragraph("3:30 - 5:30 | Thunder Client CRUD Demonstration", style="Heading 1")
    add_text(doc, "I will now demonstrate the project-specific CRUD object, TrainingClass.")
    steps = [
        "First, I sign in as an administrator. The response is 200 OK and returns a token. I add the token as a Bearer token in Thunder Client.",
        "Next, I send a POST request to create a Basic Life Support training class. The response is 201 Created and includes the new class ID.",
        "Then, I send a GET request to list all training classes. The new class appears in the response. I can also use the ID in a GET request to read one class.",
        "Next, I send a PUT request to update the class capacity and price. The response is 200 OK and shows the new values.",
        "Finally, I send a DELETE request using the class ID. The response confirms that the class was deleted. A final GET list request confirms that it no longer appears.",
    ]
    for number, step in enumerate(steps, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        number_run = paragraph.add_run(f"{number}. ")
        set_font(number_run, 11, TEAL, True)
        step_run = paragraph.add_run(step)
        set_font(step_run, 11, "000000")
    add_cue(doc, "Show", "Thunder Client Create, Read, Update, and Delete results. Hide token and password values before recording.")

    add_section(
        doc,
        "5:30 - 6:20 | User CRUD and Security Test",
        "I also tested the User API. A user can create an account, sign in, read their own profile, update profile details, and delete their own account. I tested authorization by using a student token to attempt to create a training class. The API correctly returns 403 Forbidden because only an administrator can manage class records.",
        "A user-create result and the student 403 response, with sensitive fields hidden.",
    )

    add_section(
        doc,
        "6:20 - 7:00 | Documentation and Next Release",
        "For project management, I created a Product Backlog and Task Board. I also created the first version of the External Design Document, including the project logo, database design, API test plan, and wireframes for the future landing page, class listing, authentication, and administrator management pages. In Part 2, I will connect the React frontend to these APIs, add the public landing page and navigation, and make the application visually polished. Thank you.",
        "The EDD and Product Backlog / Task Board documents if desired.",
    )

    doc.add_paragraph("Recording Checklist", style="Heading 1")
    checklist = [
        "Begin with the two-slide PowerPoint presentation.",
        "Keep the completed recording between 5 and 10 minutes.",
        "Show code and Thunder Client at a readable zoom level.",
        "Hide passwords, JWT tokens, and the MongoDB connection string.",
        "Upload the completed video to YouTube or another streaming provider and submit the link.",
    ]
    for item in checklist:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        check_run = paragraph.add_run("- ")
        set_font(check_run, 11, TEAL, True)
        item_run = paragraph.add_run(item)
        set_font(item_run, 11, "000000")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
